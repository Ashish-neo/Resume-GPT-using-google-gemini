from flask import Flask, request, jsonify
import os
import re
import PyPDF2
import docx
import google.generativeai as genai
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from io import BytesIO

# Load environment variables
load_dotenv()

# Configure the Google Generative AI API key
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

app = Flask(__name__)

def get_gemini_response(input_text, prompt):
    model = genai.GenerativeModel('gemini-1.5-flash')
    generation_config = genai.GenerationConfig(
        max_output_tokens=100,
        temperature=0.1
    )
    response = model.generate_content([input_text, prompt], generation_config=generation_config)
    return response.text

def process_file(file):
    file_name = secure_filename(file.filename)
    if file_name.lower().endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
        return text
    elif file_name.lower().endswith(".doc") or file_name.lower().endswith(".docx"):
        doc = docx.Document(file)
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
        return text
    else:
        raise ValueError("Unsupported file format. Please provide a .pdf, .doc, or .docx file.")

def process_prompt(file_prompt):
    res = []
    file_name = secure_filename(file_prompt.filename)

    if file_name.endswith('.docx'):
        doc = docx.Document(file_prompt)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = '\n'.join(full_text)
    else:
        try:
            # Handle other text-based files
            text = file_prompt.read().decode('utf-8')
        except UnicodeDecodeError:
            # Fallback for non-UTF-8 encoded files
            text = file_prompt.read().decode('latin1')  # or use 'replace' or 'ignore'

    # Split the text by periods
    text1 = text.split('.')
    for i in text1:
        res.append(i.strip())

    return res

def extract_bold_text(response_text):
    bold_text = re.findall(r'\*\*(.*?)\*\*', response_text)
    return bold_text

@app.route('/process', methods=['POST'])
def process():
    if 'file' not in request.files or 'prompt' not in request.files:
        return jsonify({"error": "No file or prompt provided"}), 400

    file = request.files['file']
    prompt_file = request.files['prompt']

    if file.filename == '' or prompt_file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    try:
        # Process the uploaded file
        file_content = process_file(file)
        
        # Process the prompt file
        prompt_text_list = process_prompt(prompt_file)
        
        responses = []
        for prompt_text in prompt_text_list:
            # Get the response from Gemini
            response_text = get_gemini_response(file_content, prompt_text)
            
            # Extract bold text if needed
            bold_text = extract_bold_text(response_text)
            
            responses.append({
                "prompt": prompt_text,
                "response": response_text
                #"bold_text": bold_text
            })

        return jsonify(responses)
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

if __name__ == '__main__':
    app.run(debug=True)
